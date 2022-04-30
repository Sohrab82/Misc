import docx

file_name = r"test.docx"

mydoc = docx.Document()
mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.add_picture(r"test.jpg",
                  width=docx.shared.Inches(5), height=docx.shared.Inches(7))
mydoc.save(file_name)


mydoc = docx.Document(file_name)
mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.save(file_name)
